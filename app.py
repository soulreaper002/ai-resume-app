import streamlit as st
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import docx
from io import BytesIO
import re

# --- Page Configuration ---
st.set_page_config(page_title="Enhanced AI Resume Builder", layout="wide")

# --- Function Definitions ---

def scrape_job_description(url):
    """Scrapes job description text from a URL."""
    try:
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
        soup = BeautifulSoup(response.text, 'html.parser')
        job_text = soup.body.get_text(separator=' ', strip=True)
        return job_text if job_text else "Could not retrieve text from URL."
    except Exception as e:
        return f"Error scraping URL: {e}"

def extract_text_from_resume(file):
    """Extracts text from an uploaded PDF or DOCX file."""
    text = ""
    try:
        if file.type == "application/pdf":
            pdf_doc = fitz.open(stream=file.read(), filetype="pdf")
            for page in pdf_doc:
                text += page.get_text()
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        st.error(f"Error reading file: {e}")
    return text

def get_gemini_response(prompt):
    """Sends a prompt to the Gemini model and returns the response."""
    # Using a more powerful model for better generation and analysis
    model = genai.GenerativeModel('gemini-1.5-pro-latest')
    response = model.generate_content(prompt)
    return response.text

def create_docx_from_text(text):
    """Creates a downloadable DOCX file from text with basic Markdown parsing."""
    doc = docx.Document()
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Prompts ---

def get_resume_prompt(content, job_desc):
    return f"""
    **Role:** You are a world-class professional resume writer and career coach specializing in ATS optimization.

    **Task:** Rewrite and tailor the provided resume content to perfectly match the given job description. Your goal is a resume that scores 90+ on an ATS scan and impresses recruiters.

    **Crucial Instruction:**
    Analyze the job description for critical skills (e.g., Python, SQL, AWS, Project Management). If the provided resume content is missing these skills, **you must strategically add them**. Integrate them into the Skills section or within the Work Experience bullet points. If adding a skill to work experience, frame it as a project contribution or a newly acquired competency. This is essential to pass the ATS filter.

    **Standard Instructions:**
    1.  **Analyze & Integrate:** Deeply analyze the **Job Description** for keywords, skills, and qualifications. Rewrite the **Resume Content** to align with it.
    2.  **Quantify Achievements:** Convert responsibilities into measurable achievements (e.g., "Increased sales by 15% by implementing a new CRM strategy").
    3.  **Format:** Use Markdown for structure. Use `**` for bolding key metrics or titles. Ensure the output is clean, professional, and contains only the resume text.
    4.  **Structure:** Follow this order: Name, Contact Info, Professional Summary, Skills, Work Experience, Education, Projects (optional).

    ---
    **Job Description:**
    ```
    {job_desc}
    ```
    ---
    **Resume Content to be Rewritten:**
    ```
    {content}
    ```
    ---
    **Required Output (Strict Markdown Format):**
    """

def get_ats_score_prompt(new_resume, job_desc):
    return f"""
    **Role:** You are an advanced Applicant Tracking System (ATS) scanner.

    **Task:** Analyze the provided resume against the job description and provide an ATS compatibility score out of 100.

    **Analysis Criteria:**
    1.  **Keyword Matching:** How well do the skills and technologies on the resume match the job description?
    2.  **Relevance of Experience:** Does the work experience align with the responsibilities listed in the job description?
    3.  **Clarity and Formatting:** Is the resume well-structured and easy for a machine to parse?

    **Output Requirement:**
    Provide a numerical score first, followed by a brief, bulleted list explaining the key reasons for the score. For example: "92/100\n- Excellent keyword alignment for 'Python' and 'Data Analysis'.\n- Strong action verbs used throughout.\n- Missing mention of 'Tableau' which was in the job description."

    ---
    **Job Description:**
    ```
    {job_desc}
    ```
    ---
    **Resume to Score:**
    ```
    {new_resume}
    ```
    ---
    """

def get_interview_questions_prompt(new_resume):
    return f"Based on this resume, generate 10 interview questions categorized as 'Easy', 'Medium', and 'Hard'.\n\n**Resume:**\n{new_resume}"

def get_resources_prompt(new_resume, job_desc):
    return f"""
    **Role:** You are a career development assistant.

    **Task:** Based on the skills mentioned in the resume and the requirements in the job description, provide a list of 3-5 high-quality, **free** online learning resources.

    **Instructions:**
    - Provide direct, **clickable URLs**.
    - Prioritize official documentation, in-depth tutorials, or full courses from reputable free platforms.
    - Format the output as a bulleted list with the resource name as the link text.

    ---
    **Job Description:**
    ```
    {job_desc}
    ```
    ---
    **Resume:**
    ```
    {new_resume}
    ```
    ---
    """

# --- Main App UI & Logic ---

st.title("🚀 Enhanced AI Resume Builder")
st.write("Generate a high-scoring, tailored resume and get interview-ready.")

try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except Exception:
    st.error("⚠️ Google API Key not found. Please add it to your Streamlit secrets.")
    st.stop()

# --- Sidebar for User Inputs ---
with st.sidebar:
    st.header("1. Job Details")
    jd_input_method = st.radio("Provide Job Description By:", ["Pasting Text", "URL"], horizontal=True)
    
    job_text = ""
    if jd_input_method == "Pasting Text":
        job_text = st.text_area("Paste Job Description Here", height=200, key="jd_text")
    else:
        job_url = st.text_input("Enter Job Posting URL", key="jd_url")

    st.header("2. Your Resume")
    resume_input_method = st.radio("Provide Your Resume By:", ["Uploading File", "Building From Scratch"], horizontal=True)

    resume_content = ""
    if resume_input_method == "Uploading File":
        uploaded_resume = st.file_uploader("Upload your resume (PDF or DOCX)", type=['pdf', 'docx'])
        if uploaded_resume:
            resume_content = extract_text_from_resume(uploaded_resume)
    else:
        with st.form(key='resume_form'):
            name = st.text_input("Full Name")
            contact = st.text_input("Email, Phone, LinkedIn URL")
            summary = st.text_area("Professional Summary")
            skills = st.text_area("Skills (comma-separated)")
            experience = st.text_area("Work Experience", height=150)
            education = st.text_area("Education", height=100)
            
            form_submit = st.form_submit_button("Save Details")
            if form_submit:
                resume_content = f"Name: {name}\nContact: {contact}\nSummary: {summary}\nSkills: {skills}\nExperience: {experience}\nEducation: {education}"
                st.success("Details saved!")

    st.header("3. Generate")
    submit_button = st.button("✨ Generate Tailored Resume & Insights", type="primary")

# --- Main Content Area for Outputs ---
col1, col2 = st.columns((2, 1))

with col1:
    st.subheader("📊 ATS Score & Analysis")
    ats_score_placeholder = st.empty()

    st.subheader("📄 Your New, Tailored Resume")
    resume_placeholder = st.container(height=600, border=True)

with col2:
    st.subheader("📋 Job Description")
    jd_placeholder = st.expander("View Job Description", expanded=True)

    st.subheader("🤔 Interview Questions")
    questions_placeholder = st.container(height=300, border=True)
    
    st.subheader("📚 Free Interview Prep Resources")
    resources_placeholder = st.container(height=300, border=True)

# --- Processing Logic ---
if submit_button:
    if jd_input_method == "URL" and job_url:
        with st.spinner("Scraping job description..."):
            job_text = scrape_job_description(job_url)
    
    if not job_text:
        st.error("Please provide a job description.")
    elif not resume_content:
        st.error("Please provide your resume details.")
    else:
        with st.spinner("AI is working its magic... 🧙‍♂️ This may take a moment."):
            # Display Job Description
            jd_placeholder.text_area("", value=job_text, height=200, disabled=True)
            
            # Generate Resume
            resume_prompt = get_resume_prompt(resume_content, job_text)
            new_resume = get_gemini_response(resume_prompt)
            resume_placeholder.markdown(new_resume)

            # Generate and Display ATS Score
            ats_prompt = get_ats_score_prompt(new_resume, job_text)
            ats_response = get_gemini_response(ats_prompt)
            
            # Extract score and feedback
            score_match = re.search(r'\d+', ats_response)
            score = int(score_match.group(0)) if score_match else 0
            feedback = ats_response
            
            with ats_score_placeholder.container():
                st.metric(label="ATS Compatibility Score", value=f"{score}/100")
                st.info(feedback, icon="💡")
            
            # Add download button inside the sidebar
            docx_file = create_docx_from_text(new_resume)
            st.sidebar.download_button(
                label="⬇️ Download Resume as DOCX",
                data=docx_file,
                file_name="Tailored_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Generate Questions
            questions_prompt = get_interview_questions_prompt(new_resume)
            questions = get_gemini_response(questions_prompt)
            questions_placeholder.markdown(questions)
            
            # Generate Resources
            resources_prompt = get_resources_prompt(new_resume, job_text)
            resources = get_gemini_response(resources_prompt)
            resources_placeholder.markdown(resources)